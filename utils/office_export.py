"""
Экспорт встреч пациента в офисные форматы.
"""

from __future__ import annotations

import re
from copy import copy, deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any

from models.db_models import (
    DOCUMENT_TYPE_MEETING,
    DOCUMENT_TYPE_PLAN,
    Document,
    Encounter,
    EncounterInformant,
    EventReportPosition,
    Patient,
    TreatmentPlanItem,
    User,
)
from utils.app_paths import get_app_base_dir, get_resource_path
from utils.word_export import WordTemplateError, render_word_template

EXCEL_PLACEHOLDER_RE = re.compile(r"{{\s*([\w.]+)\s*}}")


def safe_export_name(value: str, fallback: str = "export") -> str:
    name = re.sub(r"[^\w\-]+", "_", value.strip(), flags=re.UNICODE).strip("_")
    return name or fallback


def build_encounter_docx_filename(patient: Patient, encounter: Encounter) -> str:
    date_part = _format_date_for_filename(encounter.started_at)
    return f"pulsar_meeting_{safe_export_name(patient.callsign, 'patient')}_{date_part}.docx"


def build_plan_docx_filename(patient: Patient, plan_document: Document) -> str:
    date_part = _format_date_for_filename(plan_document.doc_date)
    number_part = safe_export_name(str(plan_document.doc_number or plan_document.id), "plan")
    patient_part = safe_export_name(patient.callsign, "patient")
    return f"pulsar_plan_{patient_part}_{number_part}_{date_part}.docx"


def build_planning_year_docx_filename(user: User, year: int) -> str:
    username = safe_export_name(user.username or user.full_name, "user")
    return f"pulsar_planning_{year}_{username}.docx"


def build_patient_encounters_xlsx_filename(patient: Patient) -> str:
    return f"pulsar_meetings_{safe_export_name(patient.callsign, 'patient')}.xlsx"


def export_patient_encounters_to_xlsx(
    patient: Patient,
    file_path: str | Path,
):
    try:
        from openpyxl import Workbook, load_workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        raise RuntimeError(
            "Для экспорта в Excel нужен пакет openpyxl. "
            "Установите зависимости из requirements.txt."
        ) from exc

    rows = collect_patient_encounter_rows(patient)
    template_path = get_excel_template_path("encounters.xlsx")
    if template_path.exists():
        _export_patient_encounters_with_xlsx_template(
            load_workbook(str(template_path)),
            patient,
            rows,
            file_path,
        )
        return

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Встречи"

    title = f"Встречи пациента: {patient.callsign or '—'}"
    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=10)
    title_cell = sheet.cell(row=1, column=1, value=title)
    title_cell.font = Font(bold=True, size=14)

    headers = [
        "Дата",
        "Пациент",
        "Личный номер",
        "Врач",
        "Результат",
        "Причина",
        "Статус",
        "Информация от пациента",
        "Документ",
        "Заметки",
    ]
    sheet.append(headers)
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    for cell in sheet[2]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(wrap_text=True, vertical="top")

    for item in rows:
        sheet.append(
            [
                item["date"],
                patient.callsign or "—",
                patient.personal_number or "—",
                item["doctor"],
                item["result"],
                item["reason"],
                item["status"],
                item["patient_info"],
                item["document_number"],
                item["notes"],
            ]
        )

    for row in sheet.iter_rows(min_row=3):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    widths = [14, 18, 16, 24, 22, 34, 16, 42, 16, 28]
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.freeze_panes = "A3"
    workbook.save(str(file_path))


def get_excel_template_path(template_name: str) -> Path:
    external_template_path = (
        get_app_base_dir() / "assets" / "templates" / "excel" / template_name
    )
    if external_template_path.exists():
        return external_template_path
    return get_resource_path("assets", "templates", "excel", template_name)


def _export_patient_encounters_with_xlsx_template(
    workbook,
    patient: Patient,
    rows: list[dict[str, Any]],
    file_path: str | Path,
):
    sheet = workbook.active
    context = {
        "title": f"Встречи пациента: {patient.callsign or '—'}",
        "patient": {
            "callsign": patient.callsign or "—",
            "personal_number": patient.personal_number or "—",
        },
    }

    row_template_index = _find_excel_row_template(sheet)
    if row_template_index:
        _render_excel_row_template(sheet, row_template_index, rows, context)

    for row in sheet.iter_rows():
        for cell in row:
            if isinstance(cell.value, str) and "{{ row." not in cell.value:
                cell.value = _render_excel_value(cell.value, context)

    workbook.save(str(file_path))


def _find_excel_row_template(sheet) -> int | None:
    for row in sheet.iter_rows():
        for cell in row:
            if isinstance(cell.value, str) and "{{ row." in cell.value:
                return cell.row
    return None


def _render_excel_row_template(sheet, row_index: int, rows: list[dict[str, Any]], context):
    output_rows = rows or [
        {
            "date": "—",
            "patient": context["patient"]["callsign"],
            "personal_number": context["patient"]["personal_number"],
            "doctor": "—",
            "result": "—",
            "reason": "—",
            "status": "—",
            "patient_info": "",
            "document_number": "—",
            "notes": "",
        }
    ]
    max_column = sheet.max_column
    template_cells = []
    for column in range(1, max_column + 1):
        source = sheet.cell(row=row_index, column=column)
        template_cells.append(
            {
                "value": source.value,
                "style": copy(source._style),
                "number_format": source.number_format,
                "font": copy(source.font),
                "fill": copy(source.fill),
                "border": copy(source.border),
                "alignment": copy(source.alignment),
                "protection": copy(source.protection),
            }
        )
    row_height = sheet.row_dimensions[row_index].height

    if len(output_rows) > 1:
        sheet.insert_rows(row_index + 1, len(output_rows) - 1)

    for offset, item in enumerate(output_rows):
        target_row = row_index + offset
        if row_height is not None:
            sheet.row_dimensions[target_row].height = row_height
        row_context = {
            **context,
            "row": {
                **item,
                "patient": context["patient"]["callsign"],
                "personal_number": context["patient"]["personal_number"],
            },
        }
        for column, template in enumerate(template_cells, start=1):
            cell = sheet.cell(row=target_row, column=column)
            cell._style = copy(template["style"])
            cell.number_format = template["number_format"]
            cell.font = copy(template["font"])
            cell.fill = copy(template["fill"])
            cell.border = copy(template["border"])
            cell.alignment = copy(template["alignment"])
            cell.protection = copy(template["protection"])
            cell.value = _render_excel_value(template["value"], row_context)


def _render_excel_value(value, context: dict[str, Any]):
    if not isinstance(value, str):
        return value

    exact_match = EXCEL_PLACEHOLDER_RE.fullmatch(value.strip())
    if exact_match:
        return _stringify_excel_value(_resolve_value(context, exact_match.group(1)))

    return EXCEL_PLACEHOLDER_RE.sub(
        lambda match: _stringify_excel_value(_resolve_value(context, match.group(1))),
        value,
    )


def _stringify_excel_value(value: Any) -> str:
    if value is None or value == "":
        return "—"
    return str(value)


def _resolve_value(context: dict[str, Any], path: str) -> Any:
    value: Any = context
    for part in path.split("."):
        if isinstance(value, dict):
            value = value.get(part)
        else:
            value = getattr(value, part, None)
        if value is None:
            return ""
    return value


def collect_patient_encounter_rows(patient: Patient) -> list[dict[str, Any]]:
    documents = Document.get_by_patient(patient.id)
    meeting_docs = [doc for doc in documents if doc.doc_type == DOCUMENT_TYPE_MEETING]
    rows = []
    for document in meeting_docs:
        encounter = document.encounter
        doctor = (
            (encounter.doctor if encounter and encounter.doctor_id else None)
            or patient.doctor
            or document.author
        )
        rows.append(
            {
                "date": _format_date(document.doc_date),
                "doctor": doctor.full_name if doctor else "—",
                "result": (
                    encounter.meeting_result_display
                    if encounter and encounter.meeting_result
                    else "—"
                ),
                "reason": (
                    encounter.reason if encounter and encounter.reason else document.summary
                )
                or "—",
                "status": encounter.status_display if encounter else "Завершен",
                "patient_info": encounter.patient_info if encounter else "",
                "document_number": _document_number(document),
                "notes": encounter.general_measures if encounter else "",
                "sort_key": document.doc_date,
            }
        )
    rows.sort(key=lambda item: _sort_value(item["sort_key"]))
    return rows


def export_encounter_to_docx(
    patient: Patient,
    encounter: Encounter,
    file_path: str | Path,
):
    document = _document_for_encounter(encounter)
    doctor = (encounter.doctor if encounter.doctor_id else None) or patient.doctor
    informants = EncounterInformant.get_by_encounter(encounter.id) if encounter.id else []

    context = {
        "patient": {
            "callsign": patient.callsign,
            "personal_number": patient.personal_number,
        },
        "encounter": {
            "word_title": _build_encounter_word_title(patient, encounter),
            "started_at": _format_datetime(encounter.started_at),
            "status": encounter.status_display,
            "meeting_result": encounter.meeting_result_display,
            "reason": encounter.reason or (document.summary if document else ""),
            "patient_info": encounter.patient_info,
            "meeting_description": encounter.meeting_description,
            "patient_tasks": encounter.patient_tasks,
            "patient_measures": encounter.patient_measures,
            "general_measures": encounter.general_measures,
        },
        "doctor": {
            "full_name": doctor.full_name if doctor else "",
        },
        "document": {
            "number": _document_number(document),
            "location": document.location if document else "",
        },
    }

    try:
        render_word_template(
            "encounter.docx",
            context,
            file_path,
            blocks={
                "informants": lambda paragraph: _insert_informants_table(
                    paragraph, informants
                ),
            },
        )
    except WordTemplateError as exc:
        raise RuntimeError(str(exc)) from exc


def export_plan_to_docx(
    patient: Patient,
    plan_document: Document,
    file_path: str | Path,
):
    if plan_document.doc_type != DOCUMENT_TYPE_PLAN:
        raise RuntimeError("Выбранный документ не является планом работы")

    items = TreatmentPlanItem.get_by_plan(plan_document.id)
    context = {
        "patient": {
            "callsign": patient.callsign or "—",
            "personal_number": patient.personal_number or "—",
        },
        "plan": {
            "number": plan_document.doc_number or f"#{plan_document.id}",
            "date": _format_date(plan_document.doc_date),
            "summary": plan_document.summary or "—",
            "items_count": len(items),
        },
        "export": {
            "created_at": _format_datetime(datetime.now()),
        },
    }

    try:
        render_word_template(
            "plan_work.docx",
            context,
            file_path,
            blocks={
                "plan_items": lambda paragraph: _insert_plan_items_table(
                    paragraph, items
                ),
            },
        )
    except WordTemplateError as exc:
        raise RuntimeError(str(exc)) from exc


def export_planning_year_to_docx(
    events: list,
    user: User,
    year: int,
    file_path: str | Path,
):
    context = {
        "planning": {
            "year": year,
            "events_count": len(events),
        },
        "user": {
            "full_name": user.full_name or user.username,
            "role": user.role_display,
        },
        "export": {
            "created_at": _format_datetime(datetime.now()),
        },
    }

    try:
        render_word_template(
            "planning_year.docx",
            context,
            file_path,
            blocks={
                "planning_events": lambda paragraph: _insert_planning_events_table(
                    paragraph, events
                ),
            },
        )
    except WordTemplateError as exc:
        raise RuntimeError(str(exc)) from exc


def _insert_plan_items_table(paragraph, items: list[TreatmentPlanItem]):
    run_properties = getattr(paragraph, "_template_run_properties", None)
    from docx.shared import Inches

    table = paragraph._parent.add_table(rows=1, cols=4, width=Inches(6.5))
    table.style = "Table Grid"
    paragraph._p.addnext(table._tbl)

    headers = ["№", "Мероприятие", "Срок исполнения", "Статус"]
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, run_properties)

    if not items:
        row = table.add_row().cells
        _set_cell_text(row[0], "—", run_properties)
        _set_cell_text(row[1], "Пункты плана не добавлены", run_properties)
        _set_cell_text(row[2], "—", run_properties)
        _set_cell_text(row[3], "—", run_properties)
        return

    for item in items:
        row = table.add_row().cells
        _set_cell_text(row[0], str(item.order_num), run_properties)
        _set_cell_text(row[1], item.event or "—", run_properties)
        _set_cell_text(row[2], _format_date(item.due_date), run_properties)
        _set_cell_text(
            row[3],
            "Выполнено" if item.is_completed else "В ожидании",
            run_properties,
        )


def _insert_planning_events_table(paragraph, events: list):
    run_properties = getattr(paragraph, "_template_run_properties", None)
    from docx.shared import Inches

    table = paragraph._parent.add_table(rows=1, cols=7, width=Inches(9.0))
    table.style = "Table Grid"
    paragraph._p.addnext(table._tbl)

    headers = [
        "Дата",
        "Время",
        "Название",
        "Тип",
        "Отделение",
        "Ответственный",
        "Отчетная позиция",
    ]
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, run_properties)

    if not events:
        row = table.add_row().cells
        _set_cell_text(row[0], "—", run_properties)
        _set_cell_text(row[1], "—", run_properties)
        _set_cell_text(row[2], "Мероприятия за выбранный год не найдены", run_properties)
        _set_cell_text(row[3], "—", run_properties)
        _set_cell_text(row[4], "—", run_properties)
        _set_cell_text(row[5], "—", run_properties)
        _set_cell_text(row[6], "—", run_properties)
        return

    for event in events:
        row = table.add_row().cells
        _set_cell_text(row[0], _format_date(event.event_date), run_properties)
        _set_cell_text(row[1], _format_event_time(event.event_time), run_properties)
        _set_cell_text(row[2], event.title or "—", run_properties)
        _set_cell_text(row[3], event.event_type_display or "—", run_properties)
        _set_cell_text(row[4], event.department_display or "—", run_properties)
        _set_cell_text(
            row[5],
            event.responsible.full_name if event.responsible else "—",
            run_properties,
        )
        _set_cell_text(row[6], _event_report_positions_text(event) or "—", run_properties)


def _build_encounter_word_title(patient: Patient, encounter: Encounter) -> str:
    patient_part = f"{patient.callsign or '—'} л.н. {patient.personal_number or '—'}"
    if encounter.meeting_result == "message":
        return f"Сообщение от {patient_part}"
    if encounter.meeting_result == "certificate":
        return f"Справка о встрече с пациентом {patient_part}"
    if encounter.meeting_result_display:
        return f"{encounter.meeting_result_display} {patient_part}"
    return f"Встреча с пациентом {patient_part}"


def _insert_informants_table(paragraph, informants: list[EncounterInformant]):
    run_properties = getattr(paragraph, "_template_run_properties", None)
    if not informants:
        _add_formatted_run(paragraph, "—", run_properties)
        return

    from docx.shared import Inches

    table = paragraph._parent.add_table(rows=1, cols=5, width=Inches(6.5))
    table.style = "Table Grid"
    paragraph._p.addnext(table._tbl)

    headers = [
        "ФИО",
        "Должность",
        "Дата рождения",
        "Место работы",
        "Суть информации / меры",
    ]
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, run_properties)

    for informant in informants:
        row = table.add_row().cells
        _set_cell_text(row[0], informant.full_name or "—", run_properties)
        _set_cell_text(row[1], informant.position or "—", run_properties)
        _set_cell_text(row[2], _format_date(informant.birth_date), run_properties)
        _set_cell_text(row[3], informant.workplace or "—", run_properties)
        _set_cell_text(
            row[4],
            _join_lines(informant.info_essence, informant.measures_taken),
            run_properties,
        )


def _set_cell_text(cell, text: str, run_properties):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _add_formatted_run(paragraph, text, run_properties)


def _add_formatted_run(paragraph, text: str, run_properties):
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, deepcopy(run_properties))
    return run


def _document_for_encounter(encounter: Encounter) -> Document | None:
    if encounter.document_id:
        return Document.get_by_id(encounter.document_id)
    return None


def _add_key_value_table(doc, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value or "—")


def _add_section(doc, title: str, text: str):
    doc.add_heading(title, level=2)
    doc.add_paragraph(text.strip() if text else "—")


def _join_lines(*values: str) -> str:
    return "\n".join(value for value in values if value) or "—"


def _document_number(document: Document | None) -> str:
    if not document:
        return "—"
    return str(document.doc_number or f"#{document.id}")


def _format_date(value) -> str:
    if not value:
        return "—"
    if hasattr(value, "strftime"):
        return value.strftime("%d.%m.%Y")
    return str(value)


def _format_datetime(value) -> str:
    if not value:
        return "—"
    if hasattr(value, "strftime"):
        return value.strftime("%d.%m.%Y %H:%M")
    return str(value)


def _format_event_time(value) -> str:
    if not value:
        return "—"
    if hasattr(value, "strftime"):
        return value.strftime("%H:%M")
    return str(value)


def _event_report_positions_text(event) -> str:
    if not getattr(event, "id", None):
        return event.report_position_display or ""

    positions = EventReportPosition.get_by_event(event.id)
    labels = [position.display_text for position in positions if position.display_text]
    if labels:
        return "\n".join(labels)
    return event.report_position_display or ""


def _format_date_for_filename(value) -> str:
    if not value:
        return "date"
    if hasattr(value, "strftime"):
        return value.strftime("%Y%m%d")
    return safe_export_name(str(value), "date")


def _sort_value(value) -> str:
    if not value:
        return ""
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value)
