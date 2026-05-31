from __future__ import annotations

import re
from copy import deepcopy
from collections.abc import Callable, Mapping
from pathlib import Path
from typing import Any

from utils.app_paths import get_app_base_dir, get_resource_path

PLACEHOLDER_RE = re.compile(r"{{\s*([\w.]+)\s*}}")
BLOCK_RE = re.compile(r"^{{\s*block:([\w.]+)\s*}}$")

BlockHandler = Callable[[Any], None]


class WordTemplateError(RuntimeError):
    pass


def render_word_template(
    template_name: str,
    context: Mapping[str, Any],
    output_path: str | Path,
    *,
    blocks: Mapping[str, BlockHandler] | None = None,
):
    try:
        from docx import Document as WordDocument
    except ImportError as exc:
        raise WordTemplateError(
            "Для экспорта в Word нужен пакет python-docx. "
            "Установите зависимости из requirements.txt."
        ) from exc

    template_path = get_word_template_path(template_name)
    if not template_path.exists():
        raise WordTemplateError(f"Не найден шаблон Word: {template_path}")

    document = WordDocument(str(template_path))
    _render_paragraphs(document.paragraphs, context, blocks or {})

    for table in document.tables:
        _render_table(table, context, blocks or {})

    document.save(str(output_path))


def get_word_template_path(template_name: str) -> Path:
    external_template_path = (
        get_app_base_dir() / "assets" / "templates" / "word" / template_name
    )
    if external_template_path.exists():
        return external_template_path
    return get_resource_path("assets", "templates", "word", template_name)


def _render_table(table, context: Mapping[str, Any], blocks: Mapping[str, BlockHandler]):
    for row in table.rows:
        for cell in row.cells:
            _render_paragraphs(cell.paragraphs, context, blocks)
            for nested_table in cell.tables:
                _render_table(nested_table, context, blocks)


def _render_paragraphs(paragraphs, context: Mapping[str, Any], blocks: Mapping[str, BlockHandler]):
    for paragraph in list(paragraphs):
        text = paragraph.text.strip()
        block_match = BLOCK_RE.match(text)
        if block_match:
            block_name = block_match.group(1)
            handler = blocks.get(block_name)
            paragraph._template_run_properties = _get_replacement_run_properties(paragraph)
            _clear_paragraph(paragraph)
            if handler:
                handler(paragraph)
            continue

        _replace_placeholders_in_paragraph(paragraph, context)


def _replace_placeholders_in_paragraph(paragraph, context: Mapping[str, Any]):
    if not PLACEHOLDER_RE.search(paragraph.text):
        return

    run_properties = _get_replacement_run_properties(paragraph)
    rendered = PLACEHOLDER_RE.sub(
        lambda match: _stringify(_resolve_value(context, match.group(1))),
        paragraph.text,
    )
    _clear_paragraph(paragraph)
    run = paragraph.add_run(rendered)
    if run_properties is not None:
        run._r.insert(0, deepcopy(run_properties))


def _get_replacement_run_properties(paragraph):
    for run in paragraph.runs:
        if PLACEHOLDER_RE.search(run.text) and run._r.rPr is not None:
            return run._r.rPr
    for run in paragraph.runs:
        if run._r.rPr is not None:
            return run._r.rPr
    return None


def _resolve_value(context: Mapping[str, Any], path: str) -> Any:
    value: Any = context
    for part in path.split("."):
        if isinstance(value, Mapping):
            value = value.get(part)
        else:
            value = getattr(value, part, None)
        if value is None:
            return ""
    return value


def _stringify(value: Any) -> str:
    if value is None or value == "":
        return "—"
    return str(value)


def _clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)
